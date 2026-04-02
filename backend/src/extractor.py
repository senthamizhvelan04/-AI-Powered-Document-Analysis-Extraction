"""
Document text extraction module for PDF, DOCX, and image files.

Provides dedicated extraction functions for each supported file format:
- PDF: Uses PyMuPDF (fitz) with block-level reading-order extraction
- DOCX: Uses python-docx for paragraphs and table cell extraction
- Image: Uses Tesseract OCR via pytesseract with Pillow preprocessing,
         with optional Google Cloud Vision API fallback
"""

import os
import logging
from io import BytesIO

import fitz  # PyMuPDF
from docx import Document
from PIL import Image, ImageEnhance, ImageFilter
import pytesseract

logger = logging.getLogger(__name__)


def extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using PyMuPDF block-level extraction.

    Opens the PDF from raw bytes and extracts text blocks from each page.
    Blocks are sorted by vertical position (y0 coordinate) to preserve
    the original reading order, which is especially important for
    multi-column layouts. Pages are separated with clear delimiters.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Extracted text as a single string with page separators.
        Returns empty string if extraction fails.
    """
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        all_text = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("blocks")

            # Sort blocks by vertical position (y0) to preserve reading order
            sorted_blocks = sorted(blocks, key=lambda b: (b[1], b[0]))

            page_text_parts = []
            for block in sorted_blocks:
                text = block[4].strip()
                if text:
                    page_text_parts.append(text)

            page_text = "\n".join(page_text_parts)
            if page_text:
                all_text.append(f"--- Page {page_num + 1} ---\n\n{page_text}")

        doc.close()

        full_text = "\n\n".join(all_text)
        logger.info(
            "PDF extraction successful: %d pages, %d characters",
            len(all_text), len(full_text)
        )
        return full_text

    except Exception as exc:
        logger.error("PDF extraction failed: %s", str(exc), exc_info=True)
        return ""


def extract_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Reads all paragraph text and all table cell text from the document.
    Paragraphs and table cells are joined with newline separators.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Extracted text as a single string.
        Returns empty string if extraction fails.
    """
    try:
        doc = Document(BytesIO(file_bytes))

        # Extract text from all paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract text from all table cells
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        table_texts.append(cell_text)

        all_text = paragraphs + table_texts
        full_text = "\n".join(all_text)

        logger.info(
            "DOCX extraction successful: %d paragraphs, %d table cells, %d characters",
            len(paragraphs), len(table_texts), len(full_text)
        )
        return full_text

    except Exception as exc:
        logger.error("DOCX extraction failed: %s", str(exc), exc_info=True)
        return ""


def extract_image(file_bytes: bytes) -> str:
    """
    Extract text from an image file using OCR.

    Primary method: Tesseract OCR via pytesseract with Pillow preprocessing.
    Preprocessing steps for improved accuracy:
      1. Convert to grayscale
      2. Enhance contrast (2x)
      3. Apply sharpening filter

    Fallback method: Google Cloud Vision API (activated when
    GOOGLE_CLOUD_VISION=true environment variable is set).

    Args:
        file_bytes: Raw bytes of the image file.

    Returns:
        Extracted text as a single string.
        Returns empty string if extraction fails.
    """
    try:
        use_cloud_vision = os.getenv("GOOGLE_CLOUD_VISION", "false").lower() == "true"

        if use_cloud_vision:
            return _extract_image_cloud_vision(file_bytes)

        # Open and preprocess the image for better OCR accuracy
        img = Image.open(BytesIO(file_bytes))
        img = img.convert("L")  # Convert to grayscale

        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)  # Boost contrast for clearer text

        img = img.filter(ImageFilter.SHARPEN)  # Sharpen edges

        # Run Tesseract OCR with PSM 6 (uniform block of text)
        text = pytesseract.image_to_string(img, lang="eng", config="--psm 6")

        logger.info(
            "Image OCR extraction successful: %d characters extracted",
            len(text.strip())
        )
        return text.strip()

    except Exception as exc:
        logger.error("Image extraction failed: %s", str(exc), exc_info=True)
        return ""


def _extract_image_cloud_vision(file_bytes: bytes) -> str:
    """
    Extract text from an image using Google Cloud Vision API.

    This is an alternative OCR engine activated by setting the
    GOOGLE_CLOUD_VISION=true environment variable. Provides higher
    accuracy for complex documents but requires a Google Cloud
    service account.

    Args:
        file_bytes: Raw bytes of the image file.

    Returns:
        Extracted text from the image.
        Returns empty string if extraction fails.
    """
    try:
        from google.cloud import vision

        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=file_bytes)
        response = client.text_detection(image=image)

        if response.error.message:
            logger.error(
                "Cloud Vision API error: %s", response.error.message
            )
            return ""

        texts = response.text_annotations
        if texts:
            full_text = texts[0].description
            logger.info(
                "Cloud Vision extraction successful: %d characters",
                len(full_text)
            )
            return full_text.strip()

        logger.warning("Cloud Vision returned no text annotations.")
        return ""

    except Exception as exc:
        logger.error(
            "Cloud Vision extraction failed: %s", str(exc), exc_info=True
        )
        return ""
